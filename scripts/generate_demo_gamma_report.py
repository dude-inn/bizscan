# -*- coding: utf-8 -*-
"""Утилита для генерации демо-отчёта Gamma из подготовленного приложения."""
from __future__ import annotations

import argparse
import sys
from pathlib import Path
from typing import Iterable, Dict, Any

BASE_DIR = Path(__file__).resolve().parent.parent
if str(BASE_DIR) not in sys.path:
    sys.path.insert(0, str(BASE_DIR))

try:
    from docx import Document  # type: ignore
except ImportError:  # pragma: no cover - подсказка пользователю
    print("❌ Не найден пакет python-docx. Установите его командой: pip install python-docx", file=sys.stderr)
    sys.exit(1)

from services.export.gamma_exporter import generate_pdf_from_report_text, create_generation
from core.logger import get_logger
from settings import (
    GAMMA_LONG_INSTRUCTIONS,
    GAMMA_COMPACT_INSTRUCTIONS,
    GAMMA_NUM_CARDS,
    GAMMA_THEME,
)

DEFAULT_INPUT = Path(r"templates/Приложение_ООО_Улыбка_Удачи_демо.docx")
DEFAULT_OUTPUT_DIR = Path("templates")
DEFAULT_COMPANY_NAME = "ООО \"Улыбка Удачи\""

log = get_logger(__name__)


def _compact_blocks(raw_text: str, max_cards: int = 55) -> str:
    blocks = [block.strip() for block in raw_text.split('\n\n') if block.strip()]
    if len(blocks) <= max_cards:
        return '\n\n'.join(blocks)

    merged: list[str] = []
    chunk: list[str] = []
    target = max(1, len(blocks) // max_cards + (1 if len(blocks) % max_cards else 0))

    for block in blocks:
        chunk.append(block)
        if len(chunk) >= target and len(merged) < max_cards - 1:
            merged.append('\n'.join(chunk))
            chunk = []
    if chunk:
        merged.append('\n'.join(chunk))

    return '\n\n'.join(merged)


def extract_text(docx_path: Path) -> str:
    """Считывает параграфы и таблицы из DOCX и превращает в текст."""
    doc = Document(docx_path)

    chunks: list[str] = []

    def extend_with(iterable: Iterable[str]) -> None:
        for item in iterable:
            item = item.strip()
            if item:
                chunks.append(item)

    extend_with(p.text for p in doc.paragraphs)

    for table in doc.tables:
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells]
            row_text = '; '.join(filter(None, cells))
            if row_text:
                chunks.append(row_text)

    if not chunks:
        raise ValueError("В исходном документе не найден текст")

    return "\n\n".join(chunks)


def prepare_report_text(docx_path: Path) -> str:
    raw = extract_text(docx_path)
    return _compact_blocks(raw)


def build_gamma_input_text(report_text: str, *, language: str = "ru") -> Dict[str, Any]:
    long_instr = (GAMMA_LONG_INSTRUCTIONS or "").strip()
    preface = f"{long_instr}\n\n---\n\n" if long_instr else ""
    sectioned = preface + report_text.replace("\n\n", "\n\n---\n\n")

    payload: Dict[str, Any] = {
        "input_text": sectioned,
        "export_as": "pdf",
        "format": "document",
        "text_mode": "generate",
        "language": language,
        "card_split": "inputTextBreaks",
        "additional_instructions": GAMMA_COMPACT_INSTRUCTIONS,
    }
    if GAMMA_NUM_CARDS and GAMMA_NUM_CARDS > 0:
        payload["num_cards"] = GAMMA_NUM_CARDS
    theme = (GAMMA_THEME or "").strip()
    if theme:
        payload["theme_name"] = theme
    return payload


def submit_generation(report_text: str, *, language: str = "ru") -> Dict[str, Any] | str:
    payload = build_gamma_input_text(report_text, language=language)
    return create_generation(**payload)


def parse_args(argv: list[str]) -> argparse.Namespace:
    parser = argparse.ArgumentParser(description=__doc__)
    parser.add_argument("--input", type=Path, default=DEFAULT_INPUT,
                        help="Путь к исходному DOCX-приложению")
    parser.add_argument("--output-dir", type=Path, default=DEFAULT_OUTPUT_DIR,
                        help="Каталог, куда сохранить PDF Gamma")
    parser.add_argument("--company-name", default=DEFAULT_COMPANY_NAME,
                        help="Название компании для именования отчёта")
    parser.add_argument("--company-inn", default=None,
                        help="ИНН компании (опционально, попадёт в имя файла)")
    parser.add_argument("--submit-only", action="store_true",
                        help="Отправить запрос в Gamma и вывести generationId без ожидания файла")
    return parser.parse_args(argv)


def main(argv: list[str]) -> int:
    args = parse_args(argv)

    if not args.input.exists():
        log.error("Файл-приложение не найден", path=str(args.input))
        print(f"❌ Файл {args.input} не найден", file=sys.stderr)
        return 1

    report_text = prepare_report_text(args.input)

    if args.submit_only:
        response = submit_generation(report_text)
        if isinstance(response, dict):
            data = response
        else:
            data = {"generationId": response}
        gen_id = data.get("generationId") or data.get("id")
        share_url = (data.get("shareUrl") or data.get("url") or data.get("publicUrl")
                     or data.get("shareURL"))
        warnings = data.get("warnings") or data.get("warning")

        print("✅ Запрос отправлен в Gamma")
        if gen_id:
            print(f"generationId: {gen_id}")
        if share_url:
            print(f"shareUrl: {share_url}")
        if warnings:
            print(f"⚠️ Предупреждения: {warnings}")
        return 0

    args.output_dir.mkdir(parents=True, exist_ok=True)

    log.info("Генерация Gamma демо-отчёта", input=str(args.input), output_dir=str(args.output_dir))
    pdf_path = generate_pdf_from_report_text(
        report_text,
        out_dir=str(args.output_dir),
        language="ru",
        company_name=args.company_name,
        company_inn=args.company_inn,
    )

    if not pdf_path:
        log.error("Gamma не вернула PDF-файл")
        print("❌ Не удалось получить PDF от Gamma. Проверьте API ключ и логи.", file=sys.stderr)
        return 2

    print(f"✅ Демо-отчёт сформирован: {pdf_path}")
    return 0


if __name__ == "__main__":
    sys.exit(main(sys.argv[1:]))
