# -*- coding: utf-8 -*-
"""
Обработчики для работы с компаниями (исправленная версия)
"""
import json
import tempfile
from typing import Optional
import asyncio
import os
from pathlib import Path
import inspect

from aiogram import Router, F
from aiogram.types import CallbackQuery, Message, InlineKeyboardMarkup, InlineKeyboardButton, BufferedInputFile
from aiogram.fsm.context import FSMContext
from aiogram.types import FSInputFile

from bot.states import SearchState, ReportState, FeedbackState
from bot.keyboards.main import choose_report_kb, report_menu_kb, choose_format_kb
from services.aggregator import fetch_company_report_markdown, fetch_company_profile
from core.logger import get_logger
from settings import FEEDBACK_CHAT_ID
from settings_texts import (
    REPORT_WAIT_HINT, GAMMA_PROGRESS_HINTS, TEXT_CHOOSE_FORMAT, TEXT_CHOOSE_FORMAT_HINT,
    TEXT_FORMAT_PDF_SELECTED, TEXT_FORMAT_PPTX_SELECTED, TEXT_CHAT_ID_ERROR,
    TEXT_REPORT_STARTING, TEXT_REPORT_PDF_FORMING, TEXT_REPORT_PPTX_FORMING,
    TEXT_REPORT_PDF_FALLBACK, TEXT_REPORT_LINK_AVAILABLE, TEXT_REPORT_ERROR_DETAILED,
    TEXT_PDF_STATUS_FORMING, TEXT_PDF_STATUS_DISABLED, TEXT_PDF_STATUS_NO_QUERY,
    TEXT_PDF_STATUS_NO_DATA, TEXT_PDF_STATUS_FORMING_DETAILED, TEXT_PDF_STATUS_ERROR,
    TEXT_PDF_STATUS_SUCCESS, TEXT_PDF_STATUS_ERROR_FINAL, TEXT_FEEDBACK_PROMPT,
    TEXT_REPORT_PAID_ORDER_MISSING, TEXT_REPORT_NO_QUERY, TEXT_REPORT_COMPANY_NOT_FOUND,
    TEXT_REPORT_SUCCESS_DOCX_ONLY, TEXT_REPORT_DOWNLOAD_WARNING, TEXT_FEEDBACK_EMPTY,
    TEXT_FEEDBACK_SUCCESS, TEXT_FEEDBACK_FAILED, TEXT_FEEDBACK_TECH_ERROR, TEXT_FEEDBACK_ADMIN_CHAT_MISSING
)

# Создаём роутер
router = Router(name="company_fixed")

@router.callback_query(F.data == "format_pdf")
async def choose_format_pdf(cb: CallbackQuery, state: FSMContext):
    await state.update_data(gamma_export_as="pdf")
    await cb.answer("Формат: PDF выбран")


@router.callback_query(F.data == "format_pptx")
async def choose_format_pptx(cb: CallbackQuery, state: FSMContext):
    await state.update_data(gamma_export_as="pptx")
    await cb.answer("Формат: PPTX выбран")


# Логгер
log = get_logger(__name__)

@router.message(F.text == "/id")
async def show_chat_id(msg: Message):
    """Показывает chat_id текущего чата (личка/группа/канал)."""
    try:
        chat_id = msg.chat.id
        chat_type = msg.chat.type
        user_id = msg.from_user.id if msg.from_user else None
        log.info("chat_id_request", chat_id=chat_id, chat_type=chat_type, user_id=user_id)
        await msg.answer(f"chat_id: {chat_id}\nchat_type: {chat_type}\nuser_id: {user_id}")
    except Exception as e:
        log.warning("chat_id_request_failed", error=str(e))
        await msg.answer("⚠️ Не удалось определить chat_id в этом чате.")

@router.callback_query(F.data.in_({"report_generate", "report_generate_pdf", "report_generate_pptx"}))
async def generate_report(cb: CallbackQuery, state: FSMContext):
    """Единый сценарий: формирование отчёта (PDF + DOCX приложение)"""
    print("DEBUG: generate_report called")  # Принудительный вывод
    log.info("generate_report: starting", user_id=cb.from_user.id)
    if cb.data == "report_generate_pdf":
        await state.update_data(gamma_export_as="pdf")
    elif cb.data == "report_generate_pptx":
        await state.update_data(gamma_export_as="pptx")
    
    # Проверяем, есть ли оплаченный заказ
    from services.orders import OrderService
    from core.config import load_settings
    settings = load_settings()
    order_service = OrderService(settings.SQLITE_PATH)
    
    data = await state.get_data()
    order_id = data.get("order_id")
    
    if order_id:
        order = await order_service.get_order(order_id)
        if not order or order["status"] != "paid":
            await cb.message.answer("❌ Заказ не оплачен. Сначала оплатите отчёт.")
            return
    else:
        # Если нет заказа, создаём бесплатный (для тестирования)
        log.warning("No order found, creating free report", user_id=cb.from_user.id)
    
    # Track report generation start
    from services.stats import StatsService
    from core.config import load_settings
    settings = load_settings()
    stats = StatsService(settings.SQLITE_PATH)
    await stats.track_event("report_start", cb.from_user.id, {"query": await state.get_data()})
    
    # Отвечаем на callback query сразу, чтобы избежать timeout
    try:
        await cb.answer()
    except Exception as e:
        log.warning("Could not answer callback query", error=str(e))
    
    # Перед генерацией убедимся, что выбран формат основного отчёта
    try:
        current = await state.get_data()
        export_as = (current.get("gamma_export_as") or "").lower()
        if export_as not in ("pdf", "pptx"):
            await cb.message.answer(
                TEXT_CHOOSE_FORMAT + "\n\n" + TEXT_CHOOSE_FORMAT_HINT,
                reply_markup=choose_format_kb(),
                disable_web_page_preview=True,
            )
            await cb.answer()
            return
    except Exception:
        pass

    # Показываем индикатор загрузки
    status_msg = await cb.message.answer(
        "⏳ Пожалуйста, подождите — идёт сбор данных и формирование отчёта.\n\n" + REPORT_WAIT_HINT
    )
    file_sent = False
    # Фоновый циклический апдейтер статуса на этапе сбора данных
    stop_cycle_event = asyncio.Event()
    async def _cycle_status_updates():
        stages = [
            "Запрашиваю данные ЕГРЮЛ и ФНС…",
            "Получаю сведения из КАД и ЕФРСБ…",
            "Анализирую госзакупки и контракты…",
            "Сверяю налоговую отчётность…",
            "Проверяю сведения о проверках и штрафах…",
            "Строю карту связей и оценку рисков…",
        ]
        frames = ["⏳", "⌛", "🔄", "🛠️", "📊", "🧮"]
        step = 0
        total = len(stages)
        while not stop_cycle_event.is_set():
            frame = frames[step % len(frames)]
            stage = stages[step % total]
            text = (
                f"{frame} {stage}"
                f"\n\nЭтап {step % total + 1} из {total}"
                f"\n\n{REPORT_WAIT_HINT}"
            )
            try:
                await status_msg.edit_text(text)
            except Exception:
                pass
            try:
                await asyncio.wait_for(stop_cycle_event.wait(), timeout=8)
            except asyncio.TimeoutError:
                step += 1
                continue
            break
    cycle_task = asyncio.create_task(_cycle_status_updates())
    
    try:
        # Получаем данные из состояния
        log.debug("get_state")
        data = await state.get_data()
        query = data.get("query", "")
        log.debug("state", query=query)
        
        # Получаем название компании и ИНН для формирования имени файла
        company_name = None
        company_inn = None
        
        # Пытаемся получить из предпросмотра
        preview = data.get("company_preview") or {}
        if isinstance(preview, dict):
            company_inn = preview.get("inn") or preview.get("ИНН")
            company_name = preview.get("name_short") or preview.get("name_full") or preview.get("name")
        
        # Если не нашли в предпросмотре, пытаемся из списка компаний
        if not company_name or not company_inn:
            companies = data.get("all_companies", []) or []
            if isinstance(companies, list):
                for c in companies:
                    if not isinstance(c, dict):
                        continue
                    c_inn = c.get("inn") or c.get("ИНН") or c.get("tax_number")
                    if str(c_inn) == str(query):
                        company_inn = c_inn
                        company_name = (c.get("НаимСокр") or c.get("name_short") or 
                                      c.get("НаимПолн") or c.get("name_full") or c.get("name"))
                        break
        
        log.debug("company_info", company_name=company_name, company_inn=company_inn)
        
        if not query:
            log.warning("No query in state", user_id=cb.from_user.id)
            await status_msg.edit_text("❌ Не указан поисковый запрос")
            return
        
        # Получаем отчёт компании через агрегатор
        log.info("fetch_report", query=query, user_id=cb.from_user.id)
        response = await fetch_company_report_markdown(query)
        log.debug("report_ready", length=len(response) if response else 0)
        
        if not response or response.startswith("❌"):
            log.warning("Invalid query or company not found", query=query, response=response[:200] if response else None, user_id=cb.from_user.id)
            await status_msg.edit_text("❌ Компания не найдена или некорректный ИНН/ОГРН")
            return
        
        log.info("report_ok", query=query, length=len(response))
        # Останавливаем цикл статусов сбора данных
        try:
            stop_cycle_event.set()
            await asyncio.sleep(0)  # отдать управление, чтобы задача завершилась
        except Exception:
            pass
        
        # Основной формат отчёта по выбору пользователя (pdf|pptx). По умолчанию PDF
        data = await state.get_data()
        export_as = (data.get("gamma_export_as") or "pdf").lower()
        main_file_path = None
        main_file_sent = False
        try:
            from settings import GAMMA_THEME
            log.info("gamma_main:start", user_id=cb.from_user.id, export_as=export_as)
            await status_msg.edit_text(
                ("⏳ Пожалуйста, подождите — идёт формирование основного PDF-отчёта.\n\n" if export_as == "pdf" else
                 "⏳ Пожалуйста, подождите — идёт формирование основной PPTX-презентации.\n\n") + REPORT_WAIT_HINT
            )
            try:
                await cb.bot.send_chat_action(cb.message.chat.id, "upload_document")
            except Exception:
                pass
            import time
            start_time = time.time()

            def update_progress(status, elapsed, timeout):
                minutes = int(elapsed // 60)
                seconds = int(elapsed % 60)
                # Добавим вращающиеся подсказки
                try:
                    hint = GAMMA_PROGRESS_HINTS[int(elapsed // 15) % len(GAMMA_PROGRESS_HINTS)]
                except Exception:
                    hint = ""
                extra = f"\n{hint}" if hint else ""
                progress_text = (
                    f"⏳ Формирую основной PDF-отчёт...{extra}\n\nСтатус: {status}\nПрошло: {minutes}м {seconds}с"
                )
                try:
                    result = status_msg.edit_text(progress_text)
                    if inspect.isawaitable(result):
                        asyncio.create_task(result)
                except Exception:
                    pass
                try:
                    send_action = cb.bot.send_chat_action(cb.message.chat.id, "upload_document")
                    if inspect.isawaitable(send_action):
                        asyncio.create_task(send_action)
                except Exception:
                    pass

            from settings import ENABLE_GAMMA_PDF
            log.info("Gamma main: checking settings", ENABLE_GAMMA_PDF=ENABLE_GAMMA_PDF, user_id=cb.from_user.id)
            main_file_path = None
            if ENABLE_GAMMA_PDF:
                log.info("Gamma main: starting generation", user_id=cb.from_user.id, export_as=export_as)
                try:
                    if export_as == "pptx":
                        from services.export.gamma_exporter import generate_pptx_from_report_text
                        main_file_path = generate_pptx_from_report_text(
                            response,
                            language="ru",
                            theme_name=GAMMA_THEME or None,
                            progress_callback=update_progress,
                            company_name=company_name,
                            company_inn=company_inn,
                        )
                    else:
                        from services.export.gamma_exporter import generate_pdf_from_report_text
                        main_file_path = generate_pdf_from_report_text(
                            response,
                            language="ru",
                            theme_name=GAMMA_THEME or None,
                            progress_callback=update_progress,
                            company_name=company_name,
                            company_inn=company_inn,
                        )
                    log.info("Gamma main: generation completed", user_id=cb.from_user.id, path=main_file_path)
                except Exception as e:
                    log.error("Gamma main: generation failed", user_id=cb.from_user.id, error=str(e))
                    main_file_path = None
            else:
                log.info("Gamma main: disabled", user_id=cb.from_user.id)
            end_time = time.time()
            log.info("gamma_main:done", user_id=cb.from_user.id, duration=end_time-start_time, path=main_file_path)
            # Если основной файл не сформировался, уведомим пользователя до отправки DOCX
            if ENABLE_GAMMA_PDF and not main_file_path:
                try:
                    await cb.message.answer(
                        "⚠️ Не удалось автоматически сформировать основной отчёт. Я отправлю приложение (DOCX)."
                    )
                except Exception:
                    pass
        except Exception as e:
            log.warning("Gamma PDF failed", error=str(e), user_id=cb.from_user.id)

        # Генерируем DOCX вместо TXT
        log.info("docx:start", user_id=cb.from_user.id)
        from docx import Document
        from docx.shared import Pt
        from docx.oxml.ns import qn
        from docx.enum.text import WD_ALIGN_PARAGRAPH

        doc = Document()
        # Базовый стиль
        style = doc.styles['Normal']
        style.font.name = 'Calibri'
        style._element.rPr.rFonts.set(qn('w:eastAsia'), 'Calibri')
        style.font.size = Pt(11)

        # Разбиваем отчёт по строкам и добавляем абзацы
        for line in response.splitlines():
            if line.strip() == '':
                doc.add_paragraph('')
                continue
            # Заголовки секций (====) делаем жирными
            if set(line.strip()) == {'='} and len(line.strip()) >= 10:
                # Это разделитель — пропускаем, т.к. предыдущая строка уже заголовок
                continue
            p = doc.add_paragraph()
            run = p.add_run(line)
            # Если предыдущая строка была заглавными буквами/заголовком
            if line.isupper() and len(line) < 60:
                run.bold = True

        # Сохраняем во временный файл
        with tempfile.NamedTemporaryFile(suffix='.docx', delete=False) as tmp:
            doc.save(tmp.name)
            temp_path = tmp.name
        log.debug("docx:saved", temp_path=temp_path)
        
        # Отправляем файл пользователю
        log.info("send_files", user_id=cb.from_user.id)
        with open(temp_path, 'rb') as file:
            # Формируем название DOCX файла
            if company_name and company_inn:
                from services.export.gamma_exporter import _safe_filename
                safe_name = _safe_filename(company_name)
                docx_filename = f"Приложение_{safe_name}_{company_inn}.docx"
            else:
                docx_filename = "company_report.docx"
            
            document = BufferedInputFile(file.read(), filename=docx_filename)
            
            # Сначала отправляем основной файл (или ссылку), затем DOCX как приложение
            if main_file_path:
                if isinstance(main_file_path, str) and main_file_path.startswith("LINK:"):
                    link = main_file_path.split("LINK:", 1)[1]
                    await cb.message.answer(
                        f"📎 {'PDF' if export_as=='pdf' else 'PPTX'}-версия доступна по ссылке: {link}"
                    )
                else:
                    with open(main_file_path, 'rb') as fpdf:
                        from aiogram.types import BufferedInputFile as BIF
                        log.debug("send_main_file", path=main_file_path, user_id=cb.from_user.id)
                        # Формируем caption
                        pdf_caption = "📄 Основной отчёт (PDF)" if export_as == "pdf" else "📊 Основной отчёт (PPTX)"
                        if company_name and company_inn:
                            from services.export.gamma_exporter import _safe_filename
                            safe_name = _safe_filename(company_name)
                            pdf_caption = (
                                f"📄 {safe_name} (ИНН: {company_inn}) - Основной отчёт (PDF)"
                                if export_as == "pdf" else
                                f"📊 {safe_name} (ИНН: {company_inn}) - Основной отчёт (PPTX)"
                            )
                        
                        await cb.message.answer_document(
                            BIF(fpdf.read(), filename=Path(main_file_path).name),
                            caption=pdf_caption
                        )
                        main_file_sent = True
            await status_msg.edit_text("✅ Отчёт готов! Отправляю приложение (DOCX)...")
            
            # Track successful report generation
            await stats.track_event("report_success", cb.from_user.id, {
                "company_name": company_name,
                "company_inn": company_inn,
                "has_pdf": main_file_sent if export_as == "pdf" else False,
                "has_docx": True
            })

            # Подсчёт сформированных отчётов и уведомление каждые 5
            try:
                await stats.track_event("gamma_generation", cb.from_user.id, {"format": export_as})
                today_cnt = await stats.get_event_count_today("gamma_generation")
                if today_cnt % 5 == 0:
                    admin_chat = str(FEEDBACK_CHAT_ID or "").strip()
                    if admin_chat:
                        await cb.bot.send_message(
                            admin_chat,
                            f"📣 Gamma отчётов сегодня: {today_cnt} (лимит 50). Пользователь #{cb.from_user.id}")
            except Exception as _e:
                log.warning("gamma_generation:notify_failed", error=str(_e))
            
            # Формируем caption для DOCX
            docx_caption = "📎 Приложение к отчёту (DOCX)"
            if company_name and company_inn:
                from services.export.gamma_exporter import _safe_filename
                safe_name = _safe_filename(company_name)
                docx_caption = f"📎 {safe_name} (ИНН: {company_inn}) - Приложение к отчёту (DOCX)"
            
            await cb.message.answer_document(
                document,
                caption=docx_caption
            )
            # Итоговое сообщение: предупреждение о скачивании и две кнопки
            from bot.keyboards.main import after_report_kb
            await cb.message.answer(
                "⚠️ Важно: обязательно скачайте файлы сейчас. Временные ссылки и кеш могут истечь, и повторная выдача потребует новой операции.\n\nВы можете оставить отзыв или вернуться в главное меню.",
                reply_markup=after_report_kb()
            )
            file_sent = True
        log.info("send_done", user_id=cb.from_user.id)
        
        # Удаляем временный файл
        Path(temp_path).unlink(missing_ok=True)
        log.debug("temp_deleted", temp_path=temp_path)
        
        # Переходим в состояние выбора типа отчёта
        await state.set_state(ReportState.CHOOSE)
        
    except Exception as e:
        log.error("Error in generate_report", error=str(e), error_type=type(e).__name__, user_id=cb.from_user.id)
        
        # Если есть оплаченный заказ, делаем возврат
        if order_id:
            try:
                from services.robokassa import RobokassaService
                from core.config import load_settings
                settings = load_settings()
                robokassa = RobokassaService(settings)
                await robokassa.refund_payment(
                    operation_id=str(order_id),
                    amount=str(settings.REPORT_PRICE),
                    reason="generation_failed"
                )
                log.info("Refund processed", order_id=order_id, user_id=cb.from_user.id)
            except Exception as refund_error:
                log.error("Refund failed", error=str(refund_error), order_id=order_id)
        
        # Если файл уже отправлен, не затираем успешный статус ошибкой
        if not file_sent:
            await status_msg.edit_text(
                "❌ Произошла ошибка при формировании отчёта.\n\n"
                "🔧 Возможные причины:\n"
                "• Временные проблемы с API\n"
                "• Недостаточно данных о компании\n"
                "• Технические неполадки\n\n"
                "💰 Если заказ был оплачен, средства будут возвращены автоматически.\n\n"
                "⏳ Попробуйте позже или обратитесь в поддержку.",
                reply_markup=InlineKeyboardMarkup(inline_keyboard=[
                    [InlineKeyboardButton(text="🔄 Повторить", callback_data="report_generate")],
                    [InlineKeyboardButton(text="🏠 Главное меню", callback_data="back_main")]
                ])
            )


@router.callback_query(F.data == "report_pdf_gamma")
async def report_pdf_gamma(cb: CallbackQuery, state: FSMContext):
    """Формирует PDF по текущему запросу"""
    await cb.answer()
    status = await cb.message.answer("⏳ Формирую PDF-отчёт…")
    try:
        from settings import ENABLE_GAMMA_PDF, GAMMA_THEME
        if not ENABLE_GAMMA_PDF:
            await status.edit_text("❌ PDF-отчёт отключён администратором")
            return
        data = await state.get_data()
        query = data.get("query", "")
        if not query:
            await status.edit_text("❌ Не указан запрос для отчёта")
            return
        
        # Получаем название компании и ИНН для формирования имени файла
        company_name = None
        company_inn = None
        
        # Пытаемся получить из предпросмотра
        preview = data.get("company_preview") or {}
        if isinstance(preview, dict):
            company_inn = preview.get("inn") or preview.get("ИНН")
            company_name = preview.get("name_short") or preview.get("name_full") or preview.get("name")
        
        # Если не нашли в предпросмотре, пытаемся из списка компаний
        if not company_name or not company_inn:
            companies = data.get("all_companies", []) or []
            if isinstance(companies, list):
                for c in companies:
                    if not isinstance(c, dict):
                        continue
                    c_inn = c.get("inn") or c.get("ИНН") or c.get("tax_number")
                    if str(c_inn) == str(query):
                        company_inn = c_inn
                        company_name = (c.get("НаимСокр") or c.get("name_short") or 
                                      c.get("НаимПолн") or c.get("name_full") or c.get("name"))
                        break
        report_text = await fetch_company_report_markdown(query)
        if not report_text or report_text.startswith("❌"):
            await status.edit_text("❌ Не удалось получить отчетные данные")
            return
        
        # Обновляем статус с информацией о времени ожидания
        await status.edit_text("⏳ Формирую PDF-версию отчёта...\n\n📄 Это может занять до 15 минут, так как нужно собрать данные из множества источников и сформировать структурированный PDF.\n\n⏰ Пожалуйста, подождите...")
        
        # Функция для обновления прогресса
        async def update_progress(status, elapsed, timeout):
            minutes = int(elapsed // 60)
            seconds = int(elapsed % 60)
            progress_text = f"⏳ Формирую PDF-версию отчёта...\n\n📄 Статус: {status}\n⏰ Прошло: {minutes}м {seconds}с\n\n⏰ Пожалуйста, подождите..."
            await status.edit_text(progress_text)
        
        from services.export.gamma_exporter import generate_pdf_from_report_text
        log.info("Starting Gamma PDF generation from button", user_id=cb.from_user.id)
        import time
        start_time = time.time()
        pdf_path = generate_pdf_from_report_text(
            report_text, 
            language="ru", 
            theme_name=GAMMA_THEME or None, 
            progress_callback=update_progress,
            company_name=company_name,
            company_inn=company_inn,
        )
        end_time = time.time()
        log.info("Gamma PDF generation from button completed", user_id=cb.from_user.id, duration=end_time-start_time, pdf_path=pdf_path)
        if not pdf_path:
            await status.edit_text("❌ Не удалось сформировать PDF-отчёт")
            return
        from aiogram.types import BufferedInputFile
        import os
        with open(pdf_path, 'rb') as f:
            await status.edit_text("✅ PDF готов!")
            await cb.message.answer_document(
                BufferedInputFile(f.read(), filename=os.path.basename(pdf_path)),
                caption="📄 PDF-версия"
            )
    except Exception as e:
        log.warning("Gamma PDF button failed", error=str(e))
        await status.edit_text("❌ Ошибка формирования PDF")


@router.callback_query(F.data == "leave_feedback")
async def leave_feedback_start(cb: CallbackQuery, state: FSMContext):
    await cb.answer()
    await state.set_state(FeedbackState.WAITING_TEXT)
    await cb.message.answer("📝 Пожалуйста, напишите ваш отзыв одним сообщением.")


@router.message(FeedbackState.WAITING_TEXT)
async def leave_feedback_collect(msg: Message, state: FSMContext):
    text = (msg.text or "").strip()
    if not text:
        await msg.answer("❗ Отзыв пустой. Напишите текст отзыва.")
        return
    admin_chat = str(FEEDBACK_CHAT_ID or "").strip()
    try:
        log.info("feedback:received", user_id=msg.from_user.id, has_text=bool(text), text_len=len(text))
        log.info("feedback:admin_chat_config", admin_chat_present=bool(admin_chat), admin_chat=admin_chat or None)
    except Exception:
        pass
    try:
        if admin_chat:
            user = msg.from_user
            header = f"📝 Отзыв от пользователя #{user.id} (@{user.username or 'нет'} | {user.full_name})"
            try:
                await msg.bot.send_message(admin_chat, header)
                log.info("feedback:header_sent", admin_chat=admin_chat)
            except Exception as e:
                log.warning("feedback:header_send_failed", error=str(e), admin_chat=admin_chat)
            # Пересылаем оригинал, чтобы сохранить вложения при необходимости
            try:
                await msg.forward(admin_chat)
                log.info("feedback:forward_ok", admin_chat=admin_chat)
            except Exception as e:
                log.warning("feedback:forward_failed", error=str(e), admin_chat=admin_chat)
        await msg.answer(
            "✅ Спасибо! Ваш отзыв отправлен.",
            reply_markup=InlineKeyboardMarkup(inline_keyboard=[
                [InlineKeyboardButton(text="🏠 Главное меню", callback_data="back_main")]
            ])
        )
    except Exception as e:
        log.warning("Feedback forward failed", error=str(e))
        await msg.answer(
            "⚠️ Не удалось отправить отзыв. Попробуйте позже.",
            reply_markup=InlineKeyboardMarkup(inline_keyboard=[
                [InlineKeyboardButton(text="🏠 Главное меню", callback_data="back_main")]
            ])
        )
    except Exception as e:
        # Доп. защита: если админ-чат не задан
        log.warning("feedback:unhandled_error", error=str(e))
        await msg.answer(
            "⚠️ Техническая ошибка при отправке отзыва. Мы уже занимаемся решением.",
            reply_markup=InlineKeyboardMarkup(inline_keyboard=[
                [InlineKeyboardButton(text="🏠 Главное меню", callback_data="back_main")]
            ])
        )
    finally:
        if not admin_chat:
            try:
                await msg.answer(
                    "ℹ️ Отзыв не отправлен администратору: не настроен FEEDBACK_CHAT_ID.",
                    reply_markup=InlineKeyboardMarkup(inline_keyboard=[
                        [InlineKeyboardButton(text="🏠 Главное меню", callback_data="back_main")]
                    ])
                )
                log.warning("feedback:admin_chat_missing")
            except Exception:
                pass
        await state.clear()


